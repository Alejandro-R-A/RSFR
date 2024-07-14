import os
from datetime import datetime, timedelta
from tkinter import messagebox

from docx import Document
from Data_base.db_connection import connect

def reemplazar_texto_delivery(doc_path, reemplazos, nombre_doc):
    doc = Document(doc_path)

    # Itera sobre todos los párrafos en el documento
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Si el texto a buscar está en el run, reemplázalo
            for texto_a_buscar, texto_nuevo in reemplazos.items():
                if texto_a_buscar in run.text:
                    run.text = run.text.replace(texto_a_buscar, str(texto_nuevo))

    # Itera sobre todas las tablas en el documento
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        # Si el texto a buscar está en el run, reemplázalo
                        for texto_a_buscar, texto_nuevo in reemplazos.items():
                            if texto_a_buscar in run.text:
                                run.text = run.text.replace(texto_a_buscar, str(texto_nuevo))

    # Crea la ruta del directorio
    dir_path = os.path.join("Documents", datetime.now().strftime("%d-%m-%Y"))

    # Crea los directorios si no existen
    os.makedirs(dir_path, exist_ok=True)

    # Crea la ruta completa del archivo
    file_path = os.path.join(dir_path, f"Certificado_entrega_{nombre_doc}.docx")

    # Guarda el archivo
    doc.save(file_path)
    messagebox.showinfo("Aviso", "Certificado de entrega generado correctamente")

def reemplazar_texto_return(doc_path, reemplazos, nombre_doc):
    doc = Document(doc_path)

    # Itera sobre todos los párrafos en el documento
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Si el texto a buscar está en el run, reemplázalo
            for texto_a_buscar, texto_nuevo in reemplazos.items():
                if texto_a_buscar in run.text:
                    run.text = run.text.replace(texto_a_buscar, str(texto_nuevo))

    # Itera sobre todas las tablas en el documento
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        # Si el texto a buscar está en el run, reemplázalo
                        for texto_a_buscar, texto_nuevo in reemplazos.items():
                            if texto_a_buscar in run.text:
                                run.text = run.text.replace(texto_a_buscar, str(texto_nuevo))

    # Crea la ruta del directorio
    dir_path = os.path.join("Documents", datetime.now().strftime("%d-%m-%Y"))

    # Crea los directorios si no existen
    os.makedirs(dir_path, exist_ok=True)

    # Crea la ruta completa del archivo
    file_path = os.path.join(dir_path, f"Certificado_devolución_{nombre_doc}.docx")

    # Guarda el archivo
    doc.save(file_path)
    messagebox.showinfo("Aviso", "Certificado de devolución generado correctamente")

def create_note_return():
    conection = connect()
    cur = conection.cursor()
    cur.execute("""
        WITH last_entry AS (
            SELECT * 
            FROM registro_entrada 
            ORDER BY id_registro_entrada DESC 
            LIMIT 1
        ), 
        persona_info AS (
            SELECT grado, ap_paterno, ap_materno, nombres 
            FROM persona 
            WHERE id_persona = (SELECT id_persona FROM instructor WHERE id_instructor = (SELECT id_instructor FROM last_entry))
        ),
        usuario_info AS (
            SELECT nombre_encargado 
            FROM usuario 
            WHERE id_usuario = 1
        )
        SELECT 
            le.fecha_entrada,
            le.id_registro_entrada,
            b.marca AS marca_bayoneta,
            a.marca AS marca_arma,
            a.modelo AS modelo_arma,
            le.motivo,
            a.num_pistola,
            a.num_cargadores,
            b.num_cuchillo,
            le.observaciones,
            pi.grado,
            pi.ap_paterno,
            pi.ap_materno,
            pi.nombres,
            ui.nombre_encargado
        FROM 
            last_entry le
            INNER JOIN instructor i ON le.id_instructor = i.id_instructor
            INNER JOIN arma a ON i.id_arma = a.id_arma
            INNER JOIN bayoneta b ON i.id_bayoneta = b.id_bayoneta
            CROSS JOIN persona_info pi
            CROSS JOIN usuario_info ui
    """)

    # Obtiene los resultados
    results = cur.fetchone()

    # Cierra la conexión
    conection.close()

    if results is not None:
        # Diccionario con los textos a buscar y los textos nuevos
        reemplazos = {
            '001': str(results[1]),
            'MAYO 21 DEL 2024': str(results[0]),
            'NORINCO': str(results[3]),
            'NP-22': str(results[4]),
            'A030835': str(results[6]),
            'DOS': str(results[7]),
            'TIPO': str(results[2]),
            'NUM': str(results[8]),
            'SIN NOVEDAD': str(results[9]),
            'Servicio de Guardia': str(results[5]),
            'Cursante': str(results[10]) + " " + str(results[11]) + " " + str(results[12]) + " " + str(results[13]),
            'Encargado': str(results[14])
        }
        return reemplazos
    else:
        print("No se encontraron resultados.")
        return {}


def create_note_delivery():
    conection = connect()
    cur = conection.cursor()
    cur.execute("""
        WITH last_entry AS (
            SELECT * 
            FROM registro_salida 
            ORDER BY id_registro_salida DESC 
            LIMIT 1
        ), 
        persona_info AS (
            SELECT grado, ap_paterno, ap_materno, nombres 
            FROM persona 
            WHERE id_persona = (SELECT id_persona FROM instructor WHERE id_instructor = (SELECT id_instructor FROM last_entry))
        ),
        usuario_info AS (
            SELECT nombre_encargado 
            FROM usuario 
            WHERE id_usuario = 1
        )
        SELECT 
            le.fecha_salida,
            le.id_registro_salida,
            b.marca AS marca_bayoneta,
            a.marca AS marca_arma,
            a.modelo AS modelo_arma,
            le.motivo,
            a.num_pistola,
            a.num_cargadores,
            b.num_cuchillo,
            le.observaciones,
            pi.grado,
            pi.ap_paterno,
            pi.ap_materno,
            pi.nombres,
            ui.nombre_encargado
        FROM 
            last_entry le
            INNER JOIN instructor i ON le.id_instructor = i.id_instructor
            INNER JOIN arma a ON i.id_arma = a.id_arma
            INNER JOIN bayoneta b ON i.id_bayoneta = b.id_bayoneta
            CROSS JOIN persona_info pi
            CROSS JOIN usuario_info ui
    """)

    # Obtiene los resultados
    results = cur.fetchone()

    # Cierra la conexión
    conection.close()

    # Diccionario con los textos a buscar y los textos nuevos
    reemplazos = {
        '001': str(results[1]),
        'MAYO 21 DEL 2024': str(results[0]),
        'NORINCO': str(results[3]),
        'NP-22': str(results[4]),
        'A030835': str(results[6]),
        'DOS': str(results[7]),
        'TIPO': str(results[2]),
        'NUM': str(results[8]),
        'SIN NOVEDAD': str(results[9]),
        'Servicio de Guardia': str(results[5]),
        'Cursante': str(results[10]) + " " + str(results[11]) + " " + str(results[12]) + " " + str(results[13]),
        'Encargado': str(results[14])
    }
    return reemplazos


# fecha_entrada = results[0]
    # id_registro_entrada = results[1]
    # marca_bayoneta = results[2]
    # marca_arma = results[3]
    # modelo_arma = results[4]
    # motivo = results[5]
    # num_pistola = results[6]
    # num_cargadores = results[7]
    # num_cuchillo = results[8]
    # observaciones = results[9]
    # grado = results[10]
    # ap_paterno = results[11]
    # ap_materno = results[12]
    # nombres = results[13]
    # nombre_encargado = results[14]