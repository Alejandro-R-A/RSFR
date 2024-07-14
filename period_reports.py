import pandas as pd
import psycopg2
import os
from tkinter import filedialog, Tk, Button, ttk, messagebox
from datetime import datetime, timedelta
from docx import Document

def center_window(root, width=300, height=200):
    # Obtiene las dimensiones de la pantalla
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()

    # Calcula la posición x e y para centrar la ventana
    x = (screen_width / 2) - (width / 2)
    y = (screen_height / 2) - (height / 2)

    root.geometry('%dx%d+%d+%d' % (width, height, x, y))


def period_reports_in():
    # Crea una ventana de Tkinter
    root = Tk()
    root.title("Selecciona el tipo de reporte")

    # Centra la ventana
    center_window(root, width=400, height=200)

    # Crea botones para cada tipo de reporte
    ttk.Button(root, text="Reporte diario", command=lambda: report_with_format_in('daily')).pack(pady=10)
    ttk.Button(root, text="Reporte semanal", command=lambda: report_with_format_in('weekly')).pack(pady=10)
    ttk.Button(root, text="Reporte mensual", command=lambda: report_with_format_in('monthly')).pack(pady=10)

    root.mainloop()  # Iniciar el mainloop para la nueva ventana

def period_reports_out():
    # Crea una ventana de Tkinter
    root = Tk()
    root.title("Selecciona el tipo de reporte")

    # Centra la ventana
    center_window(root, width=400, height=200)

    # Crea botones para cada tipo de reporte
    ttk.Button(root, text="Reporte diario", command=lambda: report_with_format_out('daily')).pack(pady=10)
    ttk.Button(root, text="Reporte semanal", command=lambda: report_with_format_out('weekly')).pack(pady=10)
    ttk.Button(root, text="Reporte mensual", command=lambda: report_with_format_out('monthly')).pack(pady=10)

    root.mainloop()  # Iniciar el mainloop para la nueva ventana

def generar_reporte(period):
    # Conecta a la base de datos
    conexion = psycopg2.connect(
        database='rsfr-db',
        user='postgres',
        password='admin',
        host='localhost',
        port="5432",
        client_encoding="UTF8"
    )

    # Crea un cursor
    cursor = conexion.cursor()

    # Obtiene la fecha actual
    now = datetime.now()

    # Define los periodos de tiempo para los reportes
    periods = {
        'daily': (now - timedelta(days=1), now),
        'weekly': (now - timedelta(weeks=1), now),
        'monthly': (now - timedelta(weeks=4), now)
    }

    # Obtiene el periodo de tiempo seleccionado
    start, end = periods[period]

    # Ejecuta la consulta SQL
    cursor.execute("""
        SELECT ROW_NUMBER () OVER (ORDER BY r.fecha_salida) AS numero,
               p.grado,
               p.ap_paterno || ' ' || p.ap_materno || ' ' || p.nombres AS nombres_apellidos,
               a.num_pistola,
               a.marca,
               a.procedencia,
               a.anio_dotacion,
               a.num_cargadores,
               r.fecha_salida,
               r.motivo,
               r.observaciones
        FROM registro_salida r
        JOIN instructor i ON r.id_instructor = i.id_instructor
        JOIN persona p ON i.id_persona = p.id_persona
        JOIN arma a ON i.id_arma = a.id_arma
        WHERE r.fecha_salida BETWEEN %s AND %s;
    """, (start, end))

    # Obtiene los resultados
    results = cursor.fetchall()

    # Crea un DataFrame con los resultados
    df = pd.DataFrame(results, columns=['numero', 'grado', 'nombres y apellidos', 'num_pistola', 'marca', 'procedencia', 'anio_dotacion', 'num_cargadores', 'fecha_salida', 'motivo', 'observaciones'])

    # Guarda el DataFrame en un archivo Excel
    df.to_excel(f'reporte_{period}.xlsx', index=False)

    # Cierra el cursor y la conexión
    cursor.close()
    conexion.close()

    print("Reporte generado exitosamente")  # Aviso de éxito


def report_with_format_in(period):
    # Conectarse a la base de datos PostgreSQL
    conn = psycopg2.connect(
        database='rsfr-db',
        user='postgres',
        password='admin',
        host='localhost',
        port="5432",
        client_encoding="UTF8"
    )

    # Obtiene la fecha actual
    now = datetime.now()

    # Define los periodos de tiempo para los reportes
    periods = {
        'daily': (now - timedelta(days=1), now),
        'weekly': (now - timedelta(weeks=1), now),
        'monthly': (now - timedelta(weeks=4), now)
    }

    # Obtiene el periodo de tiempo seleccionado
    start, end = periods[period]

    # Crear un cursor
    cur = conn.cursor()

    # Ejecutar una consulta SQL
    cur.execute("""
        SELECT ROW_NUMBER () OVER (ORDER BY r.fecha_entrada) AS numero,
            p.grado,
            p.ap_paterno || ' ' || p.ap_materno || ' ' || p.nombres AS nombres_apellidos,
            a.num_pistola,
            a.marca,
            a.procedencia,
            a.anio_dotacion,
            a.num_cargadores,
            r.fecha_entrada,
            r.motivo,
            r.observaciones
        FROM registro_entrada r
        JOIN instructor i ON r.id_instructor = i.id_instructor
        JOIN persona p ON i.id_persona = p.id_persona
        JOIN arma a ON i.id_arma = a.id_arma
        WHERE r.fecha_entrada BETWEEN %s AND %s;
    """, (start, end))

    # Obtener los resultados de la consulta
    rows = cur.fetchall()

    # Abrir el documento de Word existente
    doc = Document('format_register/format_in.docx')

    # Encontrar la tabla
    table = doc.tables[0]  # Asume que la tabla que quieres llenar es la primera tabla en el documento

    # Llenar la tabla con tus datos
    for i, row in enumerate(rows, start=1):  # Comienza en la segunda fila (índice 1) porque la primera fila contiene los títulos
        cells = table.add_row().cells  # Agrega una nueva fila al final de la tabla y obtén las celdas de esa fila
        for j, cell in enumerate(cells):
            cell.text = str(row[j])

    
    # Crea la ruta del directorio
    dir_path = os.path.join("Documents", datetime.now().strftime("%d-%m-%Y"), "reportes-entrada")

    # Crea los directorios si no existen
    os.makedirs(dir_path, exist_ok=True)

    # Crea la ruta completa del archivo
    file_path = os.path.join(dir_path, f"entrada_{period}.docx")

    # Guarda el archivo
    doc.save(file_path)
    
    # Cerrar la conexión
    cur.close()
    conn.close()
    messagebox.showinfo("Aviso", "Documento generado correctamente")


def report_with_format_out(period):
    # Conectarse a la base de datos PostgreSQL
    conn = psycopg2.connect(
        database='rsfr-db',
        user='postgres',
        password='admin',
        host='localhost',
        port="5432",
        client_encoding="UTF8"
    )

    # Obtiene la fecha actual
    now = datetime.now()

    # Define los periodos de tiempo para los reportes
    periods = {
        'daily': (now - timedelta(days=1), now),
        'weekly': (now - timedelta(weeks=1), now),
        'monthly': (now - timedelta(weeks=4), now)
    }

    # Obtiene el periodo de tiempo seleccionado
    start, end = periods[period]

    # Crear un cursor
    cur = conn.cursor()

    # Ejecutar una consulta SQL
    cur.execute("""
        SELECT ROW_NUMBER () OVER (ORDER BY r.fecha_salida) AS numero,
            p.grado,
            p.ap_paterno || ' ' || p.ap_materno || ' ' || p.nombres AS nombres_apellidos,
            a.num_pistola,
            a.marca,
            a.procedencia,
            a.anio_dotacion,
            a.num_cargadores,
            r.fecha_salida,
            r.motivo,
            r.observaciones
        FROM registro_salida r
        JOIN instructor i ON r.id_instructor = i.id_instructor
        JOIN persona p ON i.id_persona = p.id_persona
        JOIN arma a ON i.id_arma = a.id_arma
        WHERE r.fecha_salida BETWEEN %s AND %s;
    """, (start, end))

    # Obtener los resultados de la consulta
    rows = cur.fetchall()

    # Abrir el documento de Word existente
    doc = Document('format_register/format_out.docx')

    # Encontrar la tabla
    table = doc.tables[0]  # Asume que la tabla que quieres llenar es la primera tabla en el documento

    # Llenar la tabla con tus datos
    for i, row in enumerate(rows, start=1):  # Comienza en la segunda fila (índice 1) porque la primera fila contiene los títulos
        cells = table.add_row().cells  # Agrega una nueva fila al final de la tabla y obtén las celdas de esa fila
        for j, cell in enumerate(cells):
            cell.text = str(row[j])

    
    # Crea la ruta del directorio
    dir_path = os.path.join("Documents", datetime.now().strftime("%d-%m-%Y"), "reportes-salida")

    # Crea los directorios si no existen
    os.makedirs(dir_path, exist_ok=True)

    # Crea la ruta completa del archivo
    file_path = os.path.join(dir_path, f"salida_{period}.docx")

    # Guarda el archivo
    doc.save(file_path)
    
    # Cerrar la conexión
    cur.close()
    conn.close()
    messagebox.showinfo("Aviso", "Documento generado correctamente")

